"""
LLM-эксперт с YandexGPT для анализа и КОРРЕКЦИИ прогноза
Использует OpenAI-совместимый API Yandex Cloud

Особенности:
- Жёсткая привязка к веб-контексту и данным прогноза
- Улучшенное извлечение текста из HTML
- Структурированный промпт с чёткими ограничениями
"""

import os
import re
import requests
import numpy as np
from typing import Dict, List, Optional, Tuple
import json
import traceback
from datetime import datetime
from html.parser import HTMLParser
from io import BytesIO

try:
    import fitz  # PyMuPDF
    _PYMUPDF_AVAILABLE = True
except ImportError:
    _PYMUPDF_AVAILABLE = False

try:
    import docx as _docx_module
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("⚠️  openai не установлен. Установите: pip install openai")


class HTMLTextExtractor(HTMLParser):
    """
    Парсер для извлечения чистого текста из HTML.
    Удаляет скрипты, стили, комментарии и HTML-теги.
    """
    
    # Теги, содержимое которых нужно полностью игнорировать
    SKIP_TAGS = {'script', 'style', 'noscript', 'iframe', 'svg', 'canvas', 'template'}
    
    # Теги, после которых нужен перенос строки
    BLOCK_TAGS = {'p', 'div', 'br', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                  'li', 'tr', 'article', 'section', 'header', 'footer'}
    
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.skip_depth = 0  # Глубина вложенности в пропускаемые теги
        
    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.SKIP_TAGS:
            self.skip_depth += 1
        elif tag.lower() in self.BLOCK_TAGS and self.skip_depth == 0:
            self.text_parts.append('\n')
            
    def handle_endtag(self, tag):
        if tag.lower() in self.SKIP_TAGS:
            self.skip_depth = max(0, self.skip_depth - 1)
        elif tag.lower() in self.BLOCK_TAGS and self.skip_depth == 0:
            self.text_parts.append('\n')
            
    def handle_data(self, data):
        if self.skip_depth == 0:
            text = data.strip()
            if text:
                self.text_parts.append(text)
                
    def get_text(self) -> str:
        """Возвращает извлечённый текст"""
        raw_text = ' '.join(self.text_parts)
        # Убираем множественные пробелы и переносы
        clean_text = re.sub(r'\s+', ' ', raw_text)
        clean_text = re.sub(r'\n\s*\n', '\n', clean_text)
        return clean_text.strip()


class LLMExpert:
    """LLM-эксперт для анализа временных рядов и коррекции прогноза"""

    def __init__(self):
        # Читаем из переменных окружения
        self.api_key = os.getenv('YANDEX_API_KEY', '')
        self.folder_id = os.getenv('YANDEX_FOLDER_ID', '')
        self.model = os.getenv('YANDEX_MODEL', 'yandexgpt-lite')
        
        # OpenAI-совместимый API endpoint
        self.base_url = "https://ai.api.cloud.yandex.net/v1"
        
        # Инициализация OpenAI клиента
        self.client = None
        if OPENAI_AVAILABLE and self.api_key and self.folder_id:
            try:
                self.client = openai.OpenAI(
                    api_key=self.api_key,
                    base_url=self.base_url,
                    # Передаём folder_id как project (совместимость с OpenAI API)
                    default_headers={"x-folder-id": self.folder_id}
                )
                print(f"🔧 LLM Expert инициализирован (OpenAI-совместимый API)")
                print(f"   - API Key: ✓ Установлен")
                print(f"   - Folder ID: {self.folder_id}")
                print(f"   - Model: {self.model}")
            except Exception as e:
                print(f"⚠️  Ошибка инициализации OpenAI клиента: {e}")
                self.client = None
        else:
            print(f"🔧 LLM Expert инициализирован")
            print(f"   - API Key: {'✗ НЕ УСТАНОВЛЕН' if not self.api_key else '✓ Установлен'}")
            print(f"   - Folder ID: {'✗ НЕ УСТАНОВЛЕН' if not self.folder_id else self.folder_id}")
            if not OPENAI_AVAILABLE:
                print(f"   - OpenAI: ✗ Не установлен (pip install openai)")

    @staticmethod
    def extract_text_from_file(content: bytes, filename: str) -> str:
        """
        Извлекает текст и ТАБЛИЦЫ из загруженного файла (PDF или DOCX).

        Args:
            content: сырые байты файла
            filename: имя файла (используется для определения формата)

        Returns:
            Извлечённый текст или сообщение об ошибке.
        """
        fname = filename.lower()

        # ── PDF (с использованием PyMuPDF для таблиц) ────────────────
        if fname.endswith(".pdf"):
            if not _PYMUPDF_AVAILABLE:
                return "[PDF] Библиотека PyMuPDF (fitz) не установлена. pip install pymupdf"
            try:
                doc = fitz.open(stream=content, filetype="pdf")
                full_text = []
                for page in doc:
                    # Текст страницы
                    text = page.get_text()
                    if text.strip():
                        full_text.append(text)
                    
                    # Попытка извлечения таблиц
                    try:
                        tabs = page.find_tables()
                        for i, table in enumerate(tabs):
                            full_text.append(f"\n[Таблица {i+1} из PDF]:")
                            for row in table.extract():
                                row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
                                full_text.append(row_text)
                    except:
                        pass
                full = "\n".join(full_text)
                print(f"   📄 PDF '{filename}': {len(doc)} стр., {len(full)} символов")
                return full
            except Exception as e:
                return f"[PDF] Ошибка чтения '{filename}': {e}"

        # ── DOCX (с извлечением таблиц) ──────────────────────────────
        if fname.endswith(".docx") or fname.endswith(".doc"):
            if not _DOCX_AVAILABLE:
                return "[DOCX] Библиотека python-docx не установлена. pip install python-docx"
            try:
                doc = _docx_module.Document(BytesIO(content))
                full_text = []
                
                # Извлекаем параграфы
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)
                        
                # Извлекаем таблицы
                for table in doc.tables:
                    full_text.append("\n[Таблица из Word]:")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        full_text.append(row_text)
                        
                full = "\n".join(full_text)
                print(f"   📄 DOCX '{filename}': {len(full_text)} блоков, {len(full)} символов")
                return full
            except Exception as e:
                return f"[DOCX] Ошибка чтения '{filename}': {e}"

        return f"[ФАЙЛ] Неподдерживаемый формат: '{filename}'. Поддерживаются PDF и DOCX."

    def _extract_text_from_html(self, html: str) -> str:
        """
        Извлекает чистый текст из HTML, удаляя теги, скрипты и стили.
        
        Args:
            html: сырой HTML-код
            
        Returns:
            Чистый текст без HTML-разметки
        """
        try:
            parser = HTMLTextExtractor()
            parser.feed(html)
            return parser.get_text()
        except Exception as e:
            # Fallback: простое удаление тегов регулярным выражением
            text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
    
    def _extract_key_facts_regex(self, text: str, max_facts: int = 10) -> List[str]:
        """
        Regex-базовое извлечение фактов (фоллбэк, если LLM недоступен).
        Извлекает предложения с числами, датами, процентами.

        Args:
            text: исходный текст
            max_facts: максимальное количество фактов

        Returns:
            Список ключевых предложений
        """
        sentences = re.split(r'[.!?]\s+', text)
        key_facts = []
        patterns = [
            r'\d+[,.]?\d*\s*%',
            r'\d+[,.]?\d*\s*(млн|млрд|тыс|руб|\$|€|USD|RUB)',
            r'(рост|падение|снижение|увеличение|сокращение)',
            r'(прогноз|ожидается|планируется|оценивается)',
            r'(20[0-9]{2}|январ|феврал|март|апрел|май|июн|июл|август|сентябр|октябр|ноябр|декабр)',
            r'(спрос|предложение|цен[ау]|стоимость|курс)',
            r'(инфляци|ставк|ВВП|GDP|индекс)',
        ]
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20 or len(sentence) > 500:
                continue
            for pattern in patterns:
                if re.search(pattern, sentence, re.IGNORECASE):
                    key_facts.append(sentence)
                    break
            if len(key_facts) >= max_facts:
                break
        return key_facts

    def _extract_key_facts_llm(self, text: str, url: str, max_facts: int = 10) -> List[str]:
        """
        Умное LLM-базированное извлечение фактов из текста страницы.

        Передаёт очищенный текст страницы в YandexGPT и просит выделить
        конкретные числовые, экономические и трендовые факты, релевантные
        для социально-экономического прогнозирования.

        Args:
            text: очищенный текст страницы
            url: адрес источника (для лога)
            max_facts: максимальное число фактов

        Returns:
            Список фактов, извлечённых LLM, либо regex-фоллбэк
        """
        if not self.client:
            print(f"   ⚠️ LLM недоступен, используем regex-извлечение для {url}")
            return self._extract_key_facts_regex(text, max_facts)

        # Ограничиваем текст, чтобы не превысить лимит промпта
        # (_call_yandex_gpt дополнительно обрезает до 10 000 символов)
        trimmed = text[:4500] if len(text) > 4500 else text

        # Системный промпт (instructions) и пользовательский промпт (input)
        # передаются через _call_yandex_gpt, который использует проверенный
        # вызов responses.create() без /latest и с корректными параметрами
        instructions = (
            "Ты эксперт по извлечению фактов из текста для задач прогнозирования. "
            "Твоя задача — выделить конкретные числовые, трендовые и экономические факты, "
            "релевантные для социально-экономического прогнозирования временных рядов. "
            "Отвечай строго JSON-списком: {\"facts\": [\"...\", \"...\"]}"
        )
        prompt = (
            f"Из текста ниже (источник: {url}) извлеки до {max_facts} наиболее важных фактов "
            f"для прогнозирования социально-экономических показателей.\n"
            f"Требования к фактам:\n"
            f"- Конкретные числовые значения (%, руб., USD, индексы)\n"
            f"- Динамика: рост/снижение с указанием величины и периода\n"
            f"- Прогнозы или ожидания властей/аналитиков\n"
            f"- Ключевые события, влияющие на тренд\n\n"
            f"ТЕКСТ:\n{trimmed}\n\n"
            f"ОТВЕТ (строго JSON, без дополнительного текста):"
        )

        try:
            # Используем проверенный _call_yandex_gpt вместо прямого вызова API,
            # чтобы избежать ошибок формата (400 Bad Request) при дублировании
            # параметров и суффикса /latest в model URI
            raw = self._call_yandex_gpt(prompt, instructions)
            if not raw:
                raise ValueError("Пустой ответ от LLM")

            print(f"   🧠 LLM извлёк факты из {url}: {raw[:120]}...")

            # Парсинг JSON
            json_text = raw.strip()
            if "```json" in json_text:
                json_text = json_text.split("```json")[1].split("```")[0].strip()
            elif "```" in json_text:
                json_text = json_text.split("```")[1].split("```")[0].strip()

            # Ищем JSON-объект если ответ содержит лишний текст
            json_match = re.search(r'\{[^{}]*"facts"\s*:\s*\[[^\]]*\][^{}]*\}', json_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(0)

            data = json.loads(json_text)
            facts = data.get("facts", [])
            if isinstance(facts, list) and facts:
                extracted = [str(f).strip() for f in facts if str(f).strip()][:max_facts]
                print(f"   ✅ Извлечено {len(extracted)} фактов через LLM")
                return extracted

        except Exception as e:
            print(f"   ⚠️ LLM-извлечение не удалось ({e}), фоллбэк на regex")

        # Фоллбэк на regex
        return self._extract_key_facts_regex(text, max_facts)

    # Обратная совместимость: метод с прежним именем вызывает умный вариант
    def _extract_key_facts(self, text: str, max_facts: int = 10) -> List[str]:
        """Алиас для обратной совместимости. Используй _extract_key_facts_regex напрямую."""
        return self._extract_key_facts_regex(text, max_facts)

    def _fetch_web_context(self, urls: List[str], max_chars_per_url: int = 3000) -> Tuple[str, List[str]]:
        """
        Извлечение структурированного контекста из веб-ссылок.

        Для каждого URL:
        1. Загружает HTML страницы.
        2. Извлекает чистый текст с помощью HTMLTextExtractor.
        3. Вызывает LLM (_extract_key_facts_llm), чтобы та выделила
           конкретные числовые и трендовые факты, релевантные для прогнозирования.
           Если LLM недоступна — использует regex-фоллбэк.

        Args:
            urls: список URL для загрузки
            max_chars_per_url: максимум символов на один источник

        Returns:
            Tuple[полный_текст_контекста, список_ключевых_фактов]
        """
        context_parts = []
        all_key_facts = []

        for url in urls:
            try:
                response = requests.get(
                    url,
                    timeout=15,
                    headers={
                        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                        'Accept-Language': 'ru-RU,ru;q=0.9,en;q=0.8'
                    }
                )

                if response.status_code == 200:
                    # Шаг 1: Извлекаем чистый текст из HTML
                    clean_text = self._extract_text_from_html(response.text)

                    # Шаг 2: Ограничиваем длину для контекстного блока
                    display_text = clean_text[:max_chars_per_url] + '...' if len(clean_text) > max_chars_per_url else clean_text

                    # Шаг 3: Умное LLM-извлечение фактов (с фоллбэком на regex)
                    facts = self._extract_key_facts_llm(clean_text, url)
                    all_key_facts.extend(facts)

                    context_parts.append(f"\n--- ИСТОЧНИК: {url} ---\n{display_text}")
                    print(f"   ✓ Загружен: {url} ({len(clean_text)} символов, {len(facts)} фактов)")
                else:
                    context_parts.append(f"\n--- ИСТОЧНИК: {url} ---\nОшибка загрузки: HTTP {response.status_code}")
                    print(f"   ✗ Ошибка {response.status_code}: {url}")

            except requests.Timeout:
                context_parts.append(f"\n--- ИСТОЧНИК: {url} ---\nОшибка: таймаут соединения")
                print(f"   ✗ Таймаут: {url}")
            except Exception as e:
                context_parts.append(f"\n--- ИСТОЧНИК: {url} ---\nОшибка: {str(e)}")
                print(f"   ✗ Ошибка: {url} - {e}")

        full_context = "\n".join(context_parts) if context_parts else ""

        # Убираем дубликаты фактов, сохраняя порядок
        unique_facts = list(dict.fromkeys(all_key_facts))

        return full_context, unique_facts[:15]  # Максимум 15 уникальных фактов

    def test_yandex_gpt(self) -> bool:
        """Тестирование подключения к YandexGPT через OpenAI API"""
        
        if not self.client:
            print("❌ OpenAI клиент не инициализирован")
            print("💡 Проверьте:")
            print("   1. pip install openai")
            print("   2. YANDEX_API_KEY установлен")
            print("   3. YANDEX_FOLDER_ID установлен")
            return False
        
        try:
            print("🧪 Тестирую подключение к YandexGPT через OpenAI API...")
            print(f"   - Base URL: {self.base_url}")
            print(f"   - Folder ID: {self.folder_id}")
            print(f"   - Model: {self.model}")
            
            # Формируем model URI в формате Yandex Cloud
            model_uri = f"gpt://{self.folder_id}/{self.model}/latest"
            
            # Используем новый метод responses.create()
            response = self.client.responses.create(
                model=model_uri,
                temperature=0.3,
                instructions="Ты помощник для анализа данных.",
                input="Привет! Как дела?",
                max_output_tokens=100
            )
            
            print("✅ Подключение успешно!")
            print(f"   Ответ: {response.output_text[:100]}...")
            return True
            
        except Exception as e:
            print(f"❌ Тест не прошёл: {e}")
            print(f"   Тип ошибки: {type(e).__name__}")
            traceback.print_exc()
            
            # Проверка прав доступа
            if "403" in str(e) or "Permission denied" in str(e):
                print("\n💡 Ошибка 403 - проблема с доступом!")
                print("   Проверьте:")
                print("   1. Роль 'ai.languageModels.user' назначена на каталог")
                print("   2. Folder ID правильный")
                print("   3. API Key активен")
                print(f"\n📚 См. документацию: docs/YANDEX_GPT_SETUP.md")
            
            return False

    def _call_yandex_gpt(self, prompt: str, instructions: str = "") -> str:
        """Вызов YandexGPT через OpenAI-совместимый API"""
        
        if not self.client:
            print("❌ OpenAI клиент не инициализирован")
            return ""
        
        if not prompt or len(prompt.strip()) == 0:
            print("❌ Пустой промпт")
            return ""
        
        # Обрезаем промпт если слишком длинный
        max_prompt_length = 10000
        if len(prompt) > max_prompt_length:
            prompt = prompt[:max_prompt_length]
            print(f"⚠️  Промпт обрезан до {max_prompt_length} символов")
        
        try:
            print(f"📡 Отправляю запрос к YandexGPT через OpenAI API...")
            print(f"   - Folder ID: {self.folder_id}")
            print(f"   - Model: {self.model}")
            print(f"   - Input size: {len(prompt)} символов")
            
            # Формируем model URI
            model_uri = f"gpt://{self.folder_id}/{self.model}"
            
            # Используем responses.create()
            response = self.client.responses.create(
                model=model_uri,
                temperature=0.3,
                instructions=instructions or "Ты эксперт по анализу временных рядов. Отвечай ТОЛЬКО валидным JSON.",
                input=prompt,
                max_output_tokens=500
            )
            
            result_text = response.output_text
            print(f"✅ Получен ответ: {len(result_text)} символов")
            
            return result_text
            
        except Exception as e:
            error_msg = str(e)
            print(f"\n❌ YandexGPT API Error:")
            print(f"   - Error: {error_msg}")
            print(f"   - Type: {type(e).__name__}")
            
            # Детальная диагностика
            if "403" in error_msg or "Permission denied" in error_msg:
                print("\n   💡 СОВЕТ: Ошибка 403 - проблема с доступом!")
                print("   Проверьте:")
                print(f"   1. Роль 'ai.languageModels.user' назначена на каталог {self.folder_id}")
                print("   2. Folder ID правильный")
                print("   3. API Key активен")
                print(f"\n   📚 Полная инструкция: docs/YANDEX_GPT_SETUP.md")
                print(f"   🔥 Быстрое решение: YANDEX_GPT_QUICK_FIX.md")
            
            elif "401" in error_msg or "Unauthorized" in error_msg:
                print("\n   💡 СОВЕТ: Ошибка 401 - проблема с API ключом!")
                print("   Проверьте:")
                print("   1. YANDEX_API_KEY правильный")
                print("   2. API ключ активен")
                print("   3. API ключ привязан к сервисному аккаунту")
            
            elif "404" in error_msg or "Not Found" in error_msg:
                print("\n   💡 СОВЕТ: Ошибка 404 - неправильный folder ID или модель!")
                print("   Проверьте:")
                print(f"   1. YANDEX_FOLDER_ID: {self.folder_id}")
                print(f"   2. YANDEX_MODEL: {self.model}")
            
            elif "timeout" in error_msg.lower():
                print("\n   💡 СОВЕТ: Timeout!")
                print("   - Попробуйте ещё раз")
                print("   - Проверьте интернет-соединение")
            
            return ""

    def _build_system_prompt(self, has_web_context: bool) -> str:
        """
        Строит системный промпт для анализа данных.
        """
        base = """Ты эксперт-аналитик по временным рядам и социально-экономической статистике.
Твоя задача — оценить прогноз модели и выдать корректирующие коэффициенты.

ЛОГИКА КОРРЕКЦИИ:
• Коэффициент применяется как: итоговый_прогноз = прогноз_модели × коэффициент
• Если ты считаешь, что реальное значение ВЫШЕ прогноза → коэффициент > 1.0 (увеличиваем)
• Если ты считаешь, что реальное значение НИЖЕ прогноза → коэффициент < 1.0 (уменьшаем)
• Если прогноз адекватен → коэффициент = 1.0

ДИАПАЗОН КОЭФФИЦИЕНТОВ: от 0.8 до 1.2
- 0.80–0.89 = прогноз сильно завышен, нужно снизить на 11–20%
- 0.90–0.95 = прогноз умеренно завышен, снизить на 5–10%
- 0.96–0.99 = прогноз слегка завышен, снизить на 1–4%
- 1.00       = прогноз точен, коррекция не требуется
- 1.01–1.05 = прогноз слегка занижен, поднять на 1–5%
- 1.06–1.10 = прогноз умеренно занижен, поднять на 6–10%
- 1.11–1.20 = прогноз сильно занижен, поднять на 11–20%

ФОРМАТ ОТВЕТА — строго JSON, никакого другого текста:
{"weights": [список_коэффициентов], "comment": "чёткое обоснование каждого коэффициента"}"""

        if has_web_context:
            base += """

ИСПОЛЬЗОВАНИЕ ВНЕШНИХ ИСТОЧНИКОВ:
Тебе предоставлен текст реальных веб-источников. Это приоритетная информация!
Алгоритм анализа источников:
1. Прочитай текст источников внимательно
2. Определи: что они говорят о динамике показателя? (рост / стабильность / снижение)
3. Сравни вывод из источников с направлением прогноза модели
4. Если источники указывают на РОСТ, а прогноз занижен → увеличь коэффициент (>1.0)
5. Если источники указывают на СНИЖЕНИЕ, а прогноз завышен → уменьши коэффициент (<1.0)
6. В comment ОБЯЗАТЕЛЬНО укажи: что именно сказано в источниках и как это повлияло на коэффициент"""
        
        return base
    
    def _build_user_prompt(
        self,
        mean_val: float,
        std_val: float,
        trend: float,
        trend_direction: str,
        historical_data: np.ndarray,
        forecast: np.ndarray,
        lower_bound: np.ndarray,
        upper_bound: np.ndarray,
        web_context: str,
        key_facts: List[str]
    ) -> str:
        """
        Строит подробный промпт для получения осмысленных коэффициентов.
        """
        n_points = len(forecast)
        n_hist = len(historical_data)
        last_values = historical_data[-5:].tolist()
        
        # Вычисляем статистику
        last_val = float(historical_data[-1])
        first_forecast = float(forecast[0])
        last_forecast = float(forecast[-1])
        jump_pct = ((first_forecast - last_val) / last_val * 100) if last_val != 0 else 0
        total_change_pct = ((last_forecast - last_val) / last_val * 100) if last_val != 0 else 0
        
        # Анализ тренда
        expected_next = last_val + trend
        forecast_vs_trend = "выше" if first_forecast > expected_next else "ниже"
        deviation_from_trend = first_forecast - expected_next
        deviation_pct = (deviation_from_trend / expected_next * 100) if expected_next != 0 else 0
        
        # Анализ волатильности
        hist_volatility = std_val / mean_val * 100 if mean_val != 0 else 0
        
        # Формируем подробный промпт
        prompt = f"""═══════════════════════════════════════════════════════════════
ЗАДАЧА: Оценить прогноз модели и выдать корректирующие коэффициенты
═══════════════════════════════════════════════════════════════

📊 ИСТОРИЧЕСКИЕ ДАННЫЕ ({n_hist} наблюдений):
   • Последние 5 значений: {last_values}
   • Среднее: {mean_val:.2f}
   • Стд. отклонение: {std_val:.2f}
   • Волатильность: {hist_volatility:.1f}%
   • Последнее значение: {last_val:.2f}

📈 ТРЕНД:
   • Направление: {trend_direction}
   • Коэффициент: {trend:+.4f} за период
   • Ожидаемое следующее по тренду: {expected_next:.2f}

🔮 ПРОГНОЗ МОДЕЛИ ({n_points} точек):
   • Значения: {[round(f, 2) for f in forecast.tolist()]}
   • Первое значение: {first_forecast:.2f}
   • Последнее значение: {last_forecast:.2f}
   • Скачок от истории: {jump_pct:+.1f}%
   • Общее изменение: {total_change_pct:+.1f}%

⚠️ АНАЛИЗ ОТКЛОНЕНИЙ:
   • Прогноз {forecast_vs_trend} ожидаемого по тренду на {abs(deviation_pct):.1f}%
   • Отклонение: {deviation_from_trend:+.2f}"""

        # Добавляем внешний контекст если есть - ЭТО КЛЮЧЕВАЯ ЧАСТЬ
        has_context = (key_facts and len(key_facts) > 0) or (web_context and len(web_context.strip()) > 0)
        if has_context:
            prompt += f"""

🌐 ВНЕШНИЙ КОНТЕКСТ (ОБЯЗАТЕЛЬНО УЧИТЫВАТЬ!):
"""
            # Сначала выводим ключевые факты
            if key_facts and len(key_facts) > 0:
                for i, fact in enumerate(key_facts[:7], 1):
                    fact_text = fact[:200] + "..." if len(fact) > 200 else fact
                    prompt += f"   {i}. {fact_text}\n"

            # Добавляем фрагмент полного текста источников (до 1500 символов)
            if web_context and len(web_context.strip()) > 0:
                ctx_snippet = web_context.strip()[:1500]
                if len(web_context.strip()) > 1500:
                    ctx_snippet += "..."
                prompt += f"\n📄 ТЕКСТ ИСТОЧНИКОВ:\n{ctx_snippet}\n"

            prompt += """
   ⚡ ВАЖНО: Используй информацию из контекста для корректировки!
   Если источники указывают на рост/снижение - учти это в коэффициентах."""

        # Инструкции по выдаче коэффициентов
        forecast_trend_dir = "вверх" if last_forecast > first_forecast else ("вниз" if last_forecast < first_forecast else "стабильно")
        model_vs_history = "выше" if first_forecast > last_val else "ниже"
        
        prompt += f"""

═══════════════════════════════════════════════════════════════
ТВОЯ ЗАДАЧА:
═══════════════════════════════════════════════════════════════

Выдай ровно {n_points} корректирующих коэффициентов.

ШАГИ АНАЛИЗА:
1. Исторический тренд: {trend_direction} (коэф. {trend:+.4f}/период)
   → Ожидаемое следующее значение по тренду: {expected_next:.2f}
   → Прогноз модели: {first_forecast:.2f} ({model_vs_history} ожидаемого)

2. Скачок прогноза: {jump_pct:+.1f}% от последнего исторического ({last_val:.2f})
   → Реалистично ли это? Сравни с историческими колебаниями (стд={std_val:.2f}, волат.={hist_volatility:.1f}%)

3. {f'Внешние источники (КЛЮЧЕВОЙ ФАКТОР): используй текст из раздела ВНЕШНИЙ КОНТЕКСТ, определи направление (рост/снижение) и скорректируй коэффициент в нужную сторону.' if has_context else 'Внешний контекст не предоставлен — опирайся только на статистику.'}

4. Для каждой из {n_points} точек прогноза определи коэффициент.
   Прогноз идёт: {forecast_trend_dir} (от {first_forecast:.2f} до {last_forecast:.2f})

ПРАВИЛО: коэффициент > 1.0 означает ПОВЫШЕНИЕ прогноза, < 1.0 — СНИЖЕНИЕ.

ОТВЕТ (строго JSON, без дополнительного текста):
{{"weights": [{', '.join(['<0.8-1.2>' for _ in range(n_points)])}], "comment": "<что сказано в источниках, как это влияет на прогноз, почему выбраны данные коэффициенты>"}}"""
        
        return prompt

    def _extract_coefficients_from_text(self, text: str, expected_count: int) -> List[float]:
        """
        Извлекает коэффициенты из текстового ответа LLM.
        Ищет числа в диапазоне 0.8-1.2 или слова-маркеры (завышено/занижено).
        
        Args:
            text: текстовый ответ LLM
            expected_count: ожидаемое количество коэффициентов
            
        Returns:
            Список коэффициентов или пустой список
        """
        coefficients = []
        
        # 1. Пробуем найти массив чисел в квадратных скобках
        array_match = re.search(r'\[\s*([\d.,\s]+)\s*\]', text)
        if array_match:
            numbers_str = array_match.group(1)
            numbers = re.findall(r'(\d+\.?\d*)', numbers_str)
            for num in numbers:
                try:
                    val = float(num)
                    if 0.7 <= val <= 1.3:  # Расширенный диапазон для поиска
                        coefficients.append(val)
                except ValueError:
                    continue
            
            if len(coefficients) == expected_count:
                return coefficients
        
        # 2. Ищем числа типа 0.85, 1.15 в тексте
        decimal_numbers = re.findall(r'(?:^|[^\d])(\d\.\d{1,2})(?:[^\d]|$)', text)
        for num in decimal_numbers:
            try:
                val = float(num)
                if 0.8 <= val <= 1.2:  # Диапазон 0.8-1.2
                    coefficients.append(val)
            except ValueError:
                continue
        
        if len(coefficients) >= expected_count:
            return coefficients[:expected_count]
        
        # 3. Анализируем текст на слова-маркеры
        coefficients = []
        text_lower = text.lower()
        
        # Паттерны для каждой точки (расширенный диапазон 0.8-1.2)
        patterns = {
            # Сильное завышение
            'сильно завышен': 0.85,
            'значительно завышен': 0.85,
            'существенно завышен': 0.85,
            # Умеренное завышение
            'завышен': 0.92,
            'выше': 0.95,
            'переоценен': 0.92,
            # Норма
            'адекватн': 1.0,
            'норм': 1.0,
            'соответств': 1.0,
            'точн': 1.0,
            # Умеренное занижение
            'занижен': 1.08,
            'ниже': 1.05,
            'недооценен': 1.08,
            # Сильное занижение
            'сильно занижен': 1.15,
            'значительно занижен': 1.15,
            'существенно занижен': 1.15,
        }
        
        # Ищем упоминания для каждой точки
        for i in range(expected_count):
            point_markers = [f'{i+1}-', f'{i+1}.', f'{i+1})', f'перв' if i==0 else f'втор' if i==1 else f'трет' if i==2 else f'{i+1}']
            
            coef = 1.0  # По умолчанию
            for marker in point_markers:
                # Ищем контекст вокруг маркера
                marker_pos = text_lower.find(marker.lower())
                if marker_pos != -1:
                    context = text_lower[marker_pos:marker_pos+150]
                    # Сначала проверяем длинные паттерны (более специфичные)
                    for pattern, value in sorted(patterns.items(), key=lambda x: -len(x[0])):
                        if pattern in context:
                            coef = value
                            break
                    break
            
            coefficients.append(coef)
        
        # Проверяем, есть ли хоть какие-то отличия от 1.0
        if any(c != 1.0 for c in coefficients):
            return coefficients
        
        return []  # Не удалось извлечь

    def _validate_correction_factors(
        self, 
        factors: List, 
        expected_length: int,
        has_web_context: bool = False
    ) -> List[float]:
        """
        Валидирует и нормализует коэффициенты коррекции.
        
        Args:
            factors: список коэффициентов от LLM
            expected_length: ожидаемое количество коэффициентов
            has_web_context: есть ли веб-контекст (влияет на допустимый диапазон)
            
        Returns:
            Валидированный список коэффициентов
        """
        # Определяем допустимый диапазон - всегда 0.8-1.2 (±20%)
        # Расширенный диапазон позволяет учитывать внешний контекст и тренды
        min_factor, max_factor = 0.8, 1.2
        
        validated = []
        
        for i, f in enumerate(factors):
            try:
                value = float(f)
                # Ограничиваем диапазон
                value = max(min_factor, min(max_factor, value))
                validated.append(value)
            except (ValueError, TypeError):
                validated.append(1.0)  # Дефолт при ошибке
        
        # Дополняем до нужной длины если нужно
        while len(validated) < expected_length:
            validated.append(1.0)
        
        # Обрезаем если слишком много
        validated = validated[:expected_length]
        
        # Логируем если были изменения
        if factors != validated:
            print(f"⚠️  Коэффициенты скорректированы: {factors} → {validated}")
            print(f"   Допустимый диапазон: [{min_factor}, {max_factor}]")
        
        return validated

    def correct_forecast(
        self,
        historical_data: np.ndarray,
        forecast: np.ndarray,
        lower_bound: np.ndarray,
        upper_bound: np.ndarray,
        web_urls: Optional[List[str]] = None,
        extra_context: str = "",
    ) -> Dict:
        """
        Коррекция прогноза с помощью YandexGPT
        
        Args:
            historical_data: исторические значения
            forecast: прогноз модели
            lower_bound: нижняя граница 95% доверительного интервала
            upper_bound: верхняя граница 95% доверительного интервала
            web_urls: список URL для извлечения контекста
            
        Returns:
            dict с ключами:
                - corrected_forecast: скорректированный прогноз
                - corrected_lower: скорректированная нижняя граница
                - corrected_upper: скорректированная верхняя граница
                - analysis: текстовый анализ от LLM
                - reasoning: обоснование коррекции
                - confidence: уверенность модели
                - sources_used: использованные источники
                - correction_factors: применённые коэффициенты
                - correction_applied: был ли применён LLM (True/False)
        """
        
        # Если клиент не инициализирован - возвращаем базовый анализ
        if not self.client:
            return {
                'corrected_forecast': forecast,
                'corrected_lower': lower_bound,
                'corrected_upper': upper_bound,
                'analysis': self._basic_analysis(historical_data, forecast),
                'reasoning': 'LLM недоступен',
                'confidence': 0.0,
                'sources_used': [],
                'correction_factors': [1.0] * len(forecast),
                'correction_applied': False
            }
        
        # Статистика исторических данных
        mean_val = np.mean(historical_data)
        std_val = np.std(historical_data)
        trend = np.polyfit(range(len(historical_data)), historical_data, 1)[0]
        
        # Контекст из веб-источников
        web_context = ""
        key_facts = []
        if web_urls:
            print(f"📥 Загрузка веб-контекста из {len(web_urls)} источников...")
            web_context, key_facts = self._fetch_web_context(web_urls)

        # Дополнительный контекст из файлов (PDF/DOCX) — добавляем к web_context
        if extra_context and extra_context.strip():
            file_snippet = extra_context.strip()[:6000]
            if len(extra_context.strip()) > 6000:
                file_snippet += "..."
            web_context = (web_context + "\n\n--- КОНТЕКСТ ИЗ ФАЙЛОВ ---\n" + file_snippet).strip()
            # Извлекаем факты из файлового контекста тем же методом
            file_facts = self._extract_key_facts_llm(extra_context, "local_file")
            key_facts = list(dict.fromkeys(key_facts + file_facts))[:15]
            print(f"   📎 Добавлен контекст из файлов: {len(extra_context)} символов, {len(file_facts)} фактов")

        # Определяем направление тренда
        trend_direction = "растёт" if trend > 0 else "падает"
        
        # Формируем ЖЁСТКИЙ системный промпт
        instructions = self._build_system_prompt(has_web_context=bool(web_context))
        
        # Формируем основной промпт
        prompt = self._build_user_prompt(
            mean_val=mean_val,
            std_val=std_val,
            trend=trend,
            trend_direction=trend_direction,
            historical_data=historical_data,
            forecast=forecast,
            lower_bound=lower_bound,
            upper_bound=upper_bound,
            web_context=web_context,
            key_facts=key_facts
        )
        
        # Вызов YandexGPT
        llm_response = self._call_yandex_gpt(prompt, instructions)
        
        if not llm_response:
            # Fallback если LLM недоступен
            return {
                'corrected_forecast': forecast,
                'corrected_lower': lower_bound,
                'corrected_upper': upper_bound,
                'analysis': self._basic_analysis(historical_data, forecast),
                'reasoning': 'Ошибка вызова LLM',
                'confidence': 0.0,
                'sources_used': [],
                'correction_factors': [1.0] * len(forecast),
                'correction_applied': False
            }
        
        # Парсинг JSON ответа
        try:
            # Извлекаем JSON из ответа (может быть обёрнут в ```json ... ```)
            json_text = llm_response
            if "```json" in json_text:
                json_text = json_text.split("```json")[1].split("```")[0].strip()
            elif "```" in json_text:
                json_text = json_text.split("```")[1].split("```")[0].strip()
            
            # Пробуем найти JSON в ответе если он не парсится напрямую
            json_match = re.search(r'\{[^{}]*"weights"\s*:\s*\[[^\]]+\][^{}]*\}', json_text, re.DOTALL)
            if json_match:
                json_text = json_match.group(0)
            
            llm_data = json.loads(json_text)
            
            # Поддерживаем несколько форматов ответа: weights, correction_factors, coefficients
            correction_factors = (
                llm_data.get('weights') or 
                llm_data.get('correction_factors') or 
                llm_data.get('coefficients') or 
                []
            )
            analysis = llm_data.get('analysis') or llm_data.get('comment') or 'Анализ от YandexGPT'
            confidence = llm_data.get('confidence', 0.5)
            reasoning = llm_data.get('reasoning') or llm_data.get('comment') or ''
            sources_used = llm_data.get('sources_used', [])
            
            # Валидация коэффициентов
            correction_factors = self._validate_correction_factors(
                correction_factors, 
                len(forecast),
                has_web_context=bool(web_urls)
            )
            
            # Применяем коррекцию
            if len(correction_factors) == len(forecast):
                corrected_forecast = forecast * np.array(correction_factors)
                corrected_lower = lower_bound * np.array(correction_factors)
                corrected_upper = upper_bound * np.array(correction_factors)
                
                print(f"✅ Коррекция применена: {correction_factors}")
                
                # Генерируем подробный анализ
                detailed_analysis = self._generate_detailed_analysis(
                    historical_data, corrected_forecast, correction_factors, True,
                    web_context=web_context, key_facts=key_facts,
                    llm_comment=analysis
                )
                
                return {
                    'corrected_forecast': corrected_forecast,
                    'corrected_lower': corrected_lower,
                    'corrected_upper': corrected_upper,
                    'analysis': detailed_analysis,
                    'reasoning': reasoning,
                    'confidence': confidence,
                    'sources_used': sources_used,
                    'correction_factors': correction_factors,
                    'correction_applied': True
                }
            else:
                print(f"⚠️  Несоответствие длин: {len(correction_factors)} != {len(forecast)}")
                detailed_analysis = self._generate_detailed_analysis(
                    historical_data, forecast, [1.0] * len(forecast), False,
                    web_context=web_context, key_facts=key_facts,
                    llm_comment=analysis
                )
                return {
                    'corrected_forecast': forecast,
                    'corrected_lower': lower_bound,
                    'corrected_upper': upper_bound,
                    'analysis': detailed_analysis,
                    'reasoning': reasoning,
                    'confidence': confidence,
                    'sources_used': sources_used,
                    'correction_factors': [1.0] * len(forecast),
                    'correction_applied': False
                }
            
        except json.JSONDecodeError as e:
            print(f"⚠️  JSON не распознан, пробуем извлечь числа из текста...")
            print(f"   Ответ LLM: {llm_response[:300]}...")
            
            # Fallback: извлечение чисел из текстового ответа
            correction_factors = self._extract_coefficients_from_text(llm_response, len(forecast))
            
            if correction_factors and any(f != 1.0 for f in correction_factors):
                print(f"✅ Извлечены коэффициенты из текста: {correction_factors}")
                corrected_forecast = forecast * np.array(correction_factors)
                corrected_lower = lower_bound * np.array(correction_factors)
                corrected_upper = upper_bound * np.array(correction_factors)
                
                # Генерируем подробный анализ
                detailed_analysis = self._generate_detailed_analysis(
                    historical_data, corrected_forecast, correction_factors, True
                )
                
                return {
                    'corrected_forecast': corrected_forecast,
                    'corrected_lower': corrected_lower,
                    'corrected_upper': corrected_upper,
                    'analysis': detailed_analysis,
                    'reasoning': 'Коэффициенты извлечены из текстового ответа',
                    'confidence': 0.3,
                    'sources_used': [],
                    'correction_factors': correction_factors,
                    'correction_applied': True
                }
            
            # Полный fallback
            return {
                'corrected_forecast': forecast,
                'corrected_lower': lower_bound,
                'corrected_upper': upper_bound,
                'analysis': self._basic_analysis(historical_data, forecast),
                'reasoning': f'Не удалось извлечь коэффициенты из ответа LLM',
                'confidence': 0.0,
                'sources_used': [],
                'correction_factors': [1.0] * len(forecast),
                'correction_applied': False
            }

    def _basic_analysis(self, historical_data: np.ndarray, forecast: np.ndarray) -> str:
        """Базовый статистический анализ без LLM"""
        return self._generate_detailed_analysis(historical_data, forecast, [1.0] * len(forecast), False)
    
    def _generate_detailed_analysis(
        self, 
        historical_data: np.ndarray, 
        forecast: np.ndarray, 
        correction_factors: List[float],
        correction_applied: bool,
        web_context: str = "",
        key_facts: List[str] = None,
        llm_comment: str = ""
    ) -> str:
        """
        Генерирует подробный анализ прогноза с коррекцией.
        
        Включает:
        - Статистику исторических данных
        - Анализ тренда
        - Оценку прогноза
        - Применённую коррекцию
        - Информацию из веб-источников (если есть)
        """
        n = len(historical_data)
        h = len(forecast)
        
        # Статистика
        mean_hist = np.mean(historical_data)
        std_hist = np.std(historical_data)
        min_hist = np.min(historical_data)
        max_hist = np.max(historical_data)
        last_val = float(historical_data[-1])
        
        # Тренд
        if n >= 3:
            trend_coef = np.polyfit(range(n), historical_data, 1)[0]
            trend_direction = "восходящий" if trend_coef > 0 else "нисходящий"
            trend_strength = abs(trend_coef) / std_hist if std_hist > 0 else 0
            trend_desc = "сильный" if trend_strength > 0.5 else "умеренный" if trend_strength > 0.2 else "слабый"
        else:
            trend_direction = "неопределённый"
            trend_desc = ""
            trend_coef = 0
        
        # Анализ прогноза
        mean_forecast = np.mean(forecast)
        change_pct = ((mean_forecast - mean_hist) / mean_hist * 100) if mean_hist != 0 else 0
        first_jump = ((forecast[0] - last_val) / last_val * 100) if last_val != 0 else 0
        
        # Формируем текст анализа
        lines = []
        lines.append("=" * 50)
        lines.append("📊 АНАЛИЗ ПРОГНОЗА")
        lines.append("=" * 50)
        
        # Раздел 1: Исторические данные
        lines.append(f"\n📈 ИСТОРИЧЕСКИЕ ДАННЫЕ (n={n}):")
        lines.append(f"   • Период: {n} наблюдений")
        lines.append(f"   • Среднее: {mean_hist:.2f}")
        lines.append(f"   • Стд. откл.: {std_hist:.2f}")
        lines.append(f"   • Диапазон: [{min_hist:.2f}, {max_hist:.2f}]")
        lines.append(f"   • Последнее значение: {last_val:.2f}")
        
        # Раздел 2: Тренд
        lines.append(f"\n📉 ТРЕНД:")
        lines.append(f"   • Направление: {trend_desc} {trend_direction}")
        lines.append(f"   • Коэффициент: {trend_coef:+.4f} за период")
        if trend_coef != 0:
            expected_next = last_val + trend_coef
            lines.append(f"   • Ожидаемое по тренду: {expected_next:.2f}")
        
        # Раздел 3: Прогноз
        lines.append(f"\n🔮 ПРОГНОЗ (h={h}):")
        lines.append(f"   • Значения: {[round(f, 2) for f in forecast.tolist()]}")
        lines.append(f"   • Среднее прогноза: {mean_forecast:.2f}")
        lines.append(f"   • Изменение от истории: {change_pct:+.1f}%")
        lines.append(f"   • Скачок от последнего: {first_jump:+.1f}%")
        
        # Раздел 4: LLM коррекция
        lines.append(f"\n🤖 LLM-КОРРЕКЦИЯ:")
        if correction_applied and any(f != 1.0 for f in correction_factors):
            lines.append(f"   • Статус: ✅ Применена")
            lines.append(f"   • Коэффициенты: {[round(f, 3) for f in correction_factors]}")
            
            # Интерпретация коэффициентов
            interpretations = []
            for i, f in enumerate(correction_factors):
                if f < 0.99:
                    interpretations.append(f"точка {i+1}: снижение на {(1-f)*100:.1f}%")
                elif f > 1.01:
                    interpretations.append(f"точка {i+1}: повышение на {(f-1)*100:.1f}%")
            
            if interpretations:
                lines.append(f"   • Корректировки: {'; '.join(interpretations)}")
        else:
            lines.append(f"   • Статус: ⚠️ Не применена (коэффициенты = 1.0)")
            lines.append(f"   • Причина: модель считает прогноз адекватным")
        
        # Раздел 5: Вывод LLM эксперта
        if llm_comment:
            lines.append(f"\n🤖 ВЫВОД LLM-ЭКСПЕРТА:")
            # Разбиваем длинный комментарий на строки
            comment_lines = llm_comment.strip().split('. ')
            for cl in comment_lines[:6]:  # Показываем до 6 предложений
                cl = cl.strip()
                if cl:
                    lines.append(f"   • {cl}{'.' if not cl.endswith('.') else ''}")
        
        # Раздел 5б: Ключевые факты из источников
        if key_facts:
            lines.append(f"\n🌐 КЛЮЧЕВЫЕ ФАКТЫ ИЗ ИСТОЧНИКОВ:")
            for i, fact in enumerate(key_facts[:5], 1):
                fact_text = fact.strip()
                if len(fact_text) > 150:
                    fact_text = fact_text[:150] + "..."
                lines.append(f"   {i}. {fact_text}")
        
        # Раздел 6: Рекомендации
        lines.append(f"\n💡 ИНТЕРПРЕТАЦИЯ:")
        if abs(first_jump) > 10:
            lines.append(f"   ⚠️ Значительный скачок ({first_jump:+.1f}%) от последнего значения")
        
        if change_pct > 20:
            lines.append(f"   📈 Прогноз существенно выше исторического уровня (+{change_pct:.1f}%)")
        elif change_pct < -20:
            lines.append(f"   📉 Прогноз существенно ниже исторического уровня ({change_pct:.1f}%)")
        else:
            lines.append(f"   ✅ Прогноз в пределах исторической динамики (изм. {change_pct:+.1f}%)")
        
        # Итог по коррекции
        if correction_applied and any(f != 1.0 for f in correction_factors):
            avg_factor = sum(correction_factors) / len(correction_factors)
            direction = "повышен" if avg_factor > 1.0 else "снижен"
            lines.append(f"   🔧 Прогноз {direction} на {abs(avg_factor - 1) * 100:.1f}% в среднем (среднее коэф. = {avg_factor:.3f})")
        
        lines.append("=" * 50)
        
        return "\n".join(lines)


# Пример использования
if __name__ == "__main__":
    import sys
    
    print("="*60)
    print("🧪 ТЕСТ YandexGPT Expert (OpenAI API)")
    print("="*60)
    
    # Загрузка .env
    from dotenv import load_dotenv
    load_dotenv('../config/.env')
    
    # Проверка переменных окружения
    api_key = os.getenv('YANDEX_API_KEY')
    folder_id = os.getenv('YANDEX_FOLDER_ID')
    
    if not api_key or not folder_id:
        print("\n❌ Не установлены переменные окружения!")
        print("Создайте config/.env с:")
        print("  YANDEX_API_KEY=your_key")
        print("  YANDEX_FOLDER_ID=your_folder")
        sys.exit(1)
    
    # Инициализация эксперта
    expert = LLMExpert()
    
    # Тест подключения
    print("\n" + "="*60)
    print("1️⃣ Тест подключения")
    print("="*60)
    success = expert.test_yandex_gpt()
    
    if success:
        print("\n" + "="*60)
        print("2️⃣ Тест коррекции прогноза")
        print("="*60)
        
        # Тестовые данные
        historical = np.array([100, 105, 103, 108, 110, 107, 112, 115])
        forecast = np.array([118, 120, 122])
        lower = np.array([115, 117, 119])
        upper = np.array([121, 123, 125])
        
        result = expert.correct_forecast(historical, forecast, lower, upper)
        
        print(f"\n✅ Результат коррекции:")
        print(f"   Прогноз: {result['corrected_forecast']}")
        print(f"   Анализ: {result['analysis']}")
        print(f"   Обоснование: {result.get('reasoning', 'N/A')}")
        print(f"   Уверенность: {result.get('confidence', 'N/A')}")
        print(f"   Коррекция применена: {result['correction_applied']}")
    else:
        print("\n❌ Тест не прошёл")
        print("📚 См. инструкцию: docs/YANDEX_GPT_SETUP.md")
